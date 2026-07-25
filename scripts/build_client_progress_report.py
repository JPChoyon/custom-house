from pathlib import Path
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "docs" / "client" / "Custom_House_Project_Progress_Report_2026-07-23.docx"

NAVY = "17365D"
BLUE = "2E74B5"
GREEN = "1E7A46"
AMBER = "B26A00"
RED = "A61B1B"
LIGHT_BLUE = "E8EEF5"
LIGHT_GREEN = "E8F4EC"
LIGHT_AMBER = "FFF2CC"
LIGHT_RED = "FCE8E6"
LIGHT_GRAY = "F3F4F6"
TEXT = "243247"
WHITE = "FFFFFF"


def shade(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc.get_or_add_tcPr()
    tc_mar = tc.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc.append(tc_mar)
    for name, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_repeat_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    repeat = OxmlElement("w:tblHeader")
    repeat.set(qn("w:val"), "true")
    tr_pr.append(repeat)


def keep_with_next(p):
    p.paragraph_format.keep_with_next = True


def set_cell_text(cell, text, bold=False, color=TEXT, size=9):
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run(str(text))
    r.bold = bold
    r.font.name = "Calibri"
    r.font.size = Pt(size)
    r.font.color.rgb = RGBColor.from_string(color)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    margins(cell)


def add_table(doc, headers, rows, widths, status_col=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.style = "Table Grid"
    for idx, (header, width) in enumerate(zip(headers, widths)):
        table.columns[idx].width = Inches(width)
        cell = table.rows[0].cells[idx]
        cell.width = Inches(width)
        shade(cell, NAVY)
        set_cell_text(cell, header, True, WHITE, 9)
    set_repeat_header(table.rows[0])
    status_colors = {
        "DONE": (LIGHT_GREEN, GREEN),
        "MOSTLY COMPLETE": (LIGHT_GREEN, GREEN),
        "PARTIAL": (LIGHT_AMBER, AMBER),
        "VERIFY": (LIGHT_BLUE, BLUE),
        "TODO": (LIGHT_RED, RED),
        "BLOCKED": (LIGHT_RED, RED),
    }
    for row_data in rows:
        cells = table.add_row().cells
        for idx, (value, width) in enumerate(zip(row_data, widths)):
            cells[idx].width = Inches(width)
            fill, color = (WHITE, TEXT)
            if status_col == idx and value in status_colors:
                fill, color = status_colors[value]
            elif len(table.rows) % 2 == 1:
                fill = "F8FAFC"
            shade(cells[idx], fill)
            set_cell_text(cells[idx], value, idx == status_col, color, 8.5)
        table.rows[-1]._tr.get_or_add_trPr()
    doc.add_paragraph().paragraph_format.space_after = Pt(0)
    return table


def add_bullets(doc, items, level=0):
    for item in items:
        p = doc.add_paragraph(style="List Bullet" if level == 0 else "List Bullet 2")
        p.add_run(item)


def add_note(doc, title, body, fill=LIGHT_BLUE, accent=BLUE):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.columns[0].width = Inches(6.5)
    cell = table.cell(0, 0)
    shade(cell, fill)
    margins(cell, 130, 180, 130, 180)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(title)
    r.bold = True
    r.font.color.rgb = RGBColor.from_string(accent)
    r.font.size = Pt(10)
    p2 = cell.add_paragraph(body)
    p2.paragraph_format.space_after = Pt(0)
    p2.runs[0].font.size = Pt(9.5)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def page_break(doc):
    doc.add_page_break()


doc = Document()
sec = doc.sections[0]
sec.top_margin = Inches(0.78)
sec.bottom_margin = Inches(0.72)
sec.left_margin = Inches(1)
sec.right_margin = Inches(1)
sec.header_distance = Inches(0.35)
sec.footer_distance = Inches(0.35)

styles = doc.styles
normal = styles["Normal"]
normal.font.name = "Calibri"
normal.font.size = Pt(10.5)
normal.font.color.rgb = RGBColor.from_string(TEXT)
normal.paragraph_format.space_after = Pt(6)
normal.paragraph_format.line_spacing = 1.15
for name, size, color, before, after in [
    ("Title", 27, NAVY, 0, 12),
    ("Subtitle", 12, BLUE, 0, 16),
    ("Heading 1", 16, BLUE, 18, 10),
    ("Heading 2", 13, BLUE, 14, 7),
    ("Heading 3", 11.5, NAVY, 10, 5),
]:
    s = styles[name]
    s.font.name = "Calibri"
    s.font.size = Pt(size)
    s.font.bold = name != "Subtitle"
    s.font.color.rgb = RGBColor.from_string(color)
    s.paragraph_format.space_before = Pt(before)
    s.paragraph_format.space_after = Pt(after)
    s.paragraph_format.keep_with_next = True
for name in ("List Bullet", "List Bullet 2"):
    s = styles[name]
    s.font.name = "Calibri"
    s.font.size = Pt(10)
    s.paragraph_format.space_after = Pt(4)
    s.paragraph_format.left_indent = Inches(0.375 if name == "List Bullet" else 0.65)
    s.paragraph_format.first_line_indent = Inches(-0.188)

# Header/footer
header = sec.header
hp = header.paragraphs[0]
hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
hr = hp.add_run("CUSTOM HOUSE  |  CREATOR MARKETPLACE")
hr.bold = True
hr.font.size = Pt(8)
hr.font.color.rgb = RGBColor.from_string(BLUE)
footer = sec.footer
fp = footer.paragraphs[0]
fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
fr = fp.add_run("Client Progress Report  •  23 July 2026  •  Confidential")
fr.font.size = Pt(8)
fr.font.color.rgb = RGBColor(100, 110, 125)

# Cover
p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(46)
r = p.add_run("CUSTOM HOUSE")
r.bold = True
r.font.size = Pt(13)
r.font.color.rgb = RGBColor.from_string(GREEN)
p = doc.add_paragraph(style="Title")
p.add_run("Creator Marketplace\nProject Progress Report")
p = doc.add_paragraph(style="Subtitle")
p.add_run("What is live, what is complete, what remains, and the recommended path to launch")
add_note(
    doc,
    "Executive position",
    "The production foundation and core creator workflow are operational. The strongest areas are Helium-led onboarding, admin review, creator status management, dashboard access states, design submissions, and safe draft publishing. The remaining roadmap is concentrated in storefront completion, real sales/earnings analytics, creator self-service profile and collection editing, referral tracking, payouts, and final end-to-end merchant acceptance testing.",
    LIGHT_GREEN,
    GREEN,
)
add_table(
    doc,
    ["Report basis", "Current verified position"],
    [
        ("Application", "Embedded Shopify React Router + TypeScript"),
        ("Production", "Render web service + PostgreSQL + Prisma migrations"),
        ("Storefront", "Shopify theme extension, App Proxy, Helium Customer Fields, InkyBay"),
        ("Quality gate", "Lint, typecheck, 33 automated tests, and production build passed"),
        ("Roadmap estimate", "Approximately 60% of the full 10-milestone scope"),
    ],
    [1.65, 4.85],
)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(16)
r = p.add_run("Prepared for client review • Status as of 23 July 2026")
r.italic = True
r.font.size = Pt(9)
r.font.color.rgb = RGBColor(100, 110, 125)

page_break(doc)
doc.add_heading("1. Executive Summary", level=1)
doc.add_paragraph(
    "Custom House has moved beyond a prototype. It now has a deployed production application, a persistent PostgreSQL database, authenticated Shopify Admin pages, verified app-proxy customer identity, Helium-driven creator intake, creator review actions, and a storefront dashboard foundation."
)
doc.add_heading("What the client can use now", level=2)
add_bullets(doc, [
    "Applicants submit through the existing Helium Customer Fields form; no duplicate application form is required.",
    "New creator-tagged Shopify customers can synchronize into the Custom House Creator app without duplicate Creator records.",
    "Admin can review pending applications, add notes, approve, reject, suspend, or reinstate creators, with audit history and Shopify tag synchronization.",
    "Approved creators can access an authenticated storefront dashboard with profile, collection, submissions, and overview sections.",
    "Creators can submit saved InkyBay design URLs; admin publishing is draft-first, retry-safe, and designed to prevent duplicate or partially published products.",
    "The live theme header includes responsive account, dashboard, application, login, and logout navigation beside the cart.",
])
doc.add_heading("Important limitations still visible", level=2)
add_bullets(doc, [
    "Dashboard sales, earnings, and order totals are intentionally marked as not configured because order access and commission rules are not yet approved.",
    "Helium’s reserved metafield namespace prevents full server-side reading of every custom answer unless fields are moved to an app-owned namespace or Helium Advanced API access is supplied.",
    "Creator self-service profile editing, collection editing, referral tracking, commission management, earnings, and payouts are not complete.",
    "Full storefront, InkyBay customization, cart, checkout, order, refund, and mobile acceptance testing still requires merchant-led live verification.",
])
add_note(
    doc,
    "How to read the status",
    "DONE means implemented and verified in the codebase. MOSTLY COMPLETE means the main workflow exists with limited merchant verification or a small data gap. PARTIAL means meaningful work exists but the milestone is not client-complete. VERIFY means a merchant-controlled Shopify, theme, publication, or third-party setting must be checked. TODO means not yet implemented.",
)

doc.add_heading("2. Milestone Status at a Glance", level=1)
matrix = [
    ("1", "Storefront Foundation", "PARTIAL", "Header/account navigation and creator pages live; full storefront checklist and final theme QA remain."),
    ("2", "InkyBay Integration", "PARTIAL", "Installed and preserved; saved-design and buy-only compatibility exist; complete DesignLab/cart/checkout QA remains."),
    ("3", "Creator Onboarding", "MOSTLY COMPLETE", "Helium is the only form; sync, statuses, admin review, notes, history, validation, and audit logs implemented."),
    ("4", "Creator Dashboard", "PARTIAL", "Access states and responsive overview are live; real sales, earnings, orders, and top-selling data remain."),
    ("5", "Creator Profile", "PARTIAL", "Names, bio, photo, portfolio, upload and storefront display exist; full self-service editing remains."),
    ("6", "Collection Management", "PARTIAL", "Approval can create/publish creator collections; creator editing and collection lifecycle UI remain."),
    ("7", "Product Creation", "PARTIAL", "Submission, base product reference, saved InkyBay URL and approval flow exist; complete live UX QA remains."),
    ("8", "Automated Publishing", "MOSTLY COMPLETE", "Draft-first, idempotent publishing and retry safety implemented; production transaction QA remains."),
    ("9", "Referral System", "TODO", "Referral links, codes, clicks, attributed orders and revenue reporting are not implemented."),
    ("10", "Admin Management", "PARTIAL", "Creators, applications, submissions, notes, audit and setup tools exist; commissions, earnings, payouts and final QA remain."),
]
add_table(doc, ["#", "Milestone", "Status", "Evidence / remaining"], matrix, [0.35, 1.35, 1.05, 3.75], 2)
add_note(
    doc,
    "Progress estimate",
    "The full ten-milestone roadmap is approximately 60% complete. This is an engineering readiness estimate—not a billing or contractual completion percentage. Milestones 3 and 8 are closest to completion; Milestone 9 has not started.",
    LIGHT_AMBER,
    AMBER,
)

page_break(doc)
doc.add_heading("3. Completed Production Foundation", level=1)
add_table(
    doc,
    ["Area", "Verified implementation"],
    [
        ("Hosting", "Production application deployed on Render with a public health endpoint."),
        ("Database", "Prisma uses PostgreSQL; Session, ShopConfig, Creator, CreatorApplication, DesignSubmission and AuditLog persist in production."),
        ("Migrations", "Reviewed PostgreSQL baseline and subsequent creator/Helium migrations are present; production uses prisma migrate deploy."),
        ("Authentication", "Embedded admin loaders, webhooks and app-proxy endpoints follow Shopify React Router authentication patterns."),
        ("Identity safety", "Storefront customer identity is taken only from Shopify-verified logged_in_customer_id; browser-provided customer IDs are not trusted."),
        ("Privacy", "No tokens, cookies, database URLs or raw GraphQL errors are exposed in the embedded or storefront UI."),
        ("Scopes", "Product, customer, metaobject, file and publication read/write scopes are preserved."),
        ("Health", "GET /health returns only {\"status\":\"ok\"}."),
    ],
    [1.35, 5.15],
)
doc.add_heading("Shopify integrations already wired", level=2)
add_bullets(doc, [
    "Customer create/update webhooks for Helium/Flow synchronization.",
    "Product and collection delete webhooks plus Shopify compliance webhooks.",
    "Production App Proxy at /apps/customhouse routed to the Render application.",
    "Theme app extension components: Creator Dashboard, Creator Submission, Creator Attribution, Buy-only Product Form, and InkyBay Compatibility Embed.",
    "Managed embedded installation and authenticated Admin GraphQL usage.",
])

doc.add_heading("4. Creator Journey — Current Workflow", level=1)
add_table(
    doc,
    ["Stage", "Current behavior", "Status"],
    [
        ("Apply", "Customer submits the Helium Customer Fields form.", "DONE"),
        ("Tag", "Helium/Shopify Flow assigns creator-applicant and creator-pending.", "VERIFY"),
        ("Synchronize", "Customer webhook upserts one Creator per shop + Shopify customer GID.", "DONE"),
        ("Review", "Application appears in the pending Applications view with profile/context fields where accessible.", "DONE"),
        ("Decision", "Approve, Reject, Suspend and Reinstate synchronize app status, tags and audit history.", "DONE"),
        ("Collection", "Approval creates the creator collection if missing and preserves customer ownership mapping.", "MOSTLY COMPLETE"),
        ("Dashboard", "Verified customer sees the correct logged-out/not-applied/pending/approved/rejected/suspended state.", "DONE"),
        ("Submit design", "Creator saves a design submission linked to an allowed base product and InkyBay saved-design URL.", "DONE"),
        ("Publish", "Admin approval creates/configures a draft creator product, assigns collection and publishes only after the transaction succeeds.", "MOSTLY COMPLETE"),
    ],
    [1.1, 4.25, 1.15],
    2,
)

page_break(doc)
doc.add_heading("5. Detailed Milestone Review", level=1)
details = [
    ("Milestone 1 — Storefront Foundation", "PARTIAL",
     ["Live responsive header profile menu placed beside cart.", "Account, Creator Dashboard, Become a Creator, login and logout links are present.", "Creator dashboard and application destination pages return successfully."],
     ["Finish/verify announcement bar, footer, homepage sections, collections, product templates, cart, checkout, content pages, policies, sample products, broken links and full mobile QA."]),
    ("Milestone 2 — InkyBay Integration", "PARTIAL",
     ["Existing InkyBay installation and merchant workflow preserved.", "Manual saved-design URL is the supported provider path.", "Buy-only product form and compatibility extension are included.", "Publishing logic does not invent or call undocumented InkyBay APIs."],
     ["Merchant must verify DesignLab settings, customization fields, pricing behavior, cart display, checkout and test orders on the live theme."]),
    ("Milestone 3 — Creator Onboarding", "MOSTLY COMPLETE",
     ["React Router/TypeScript backend and PostgreSQL data models.", "Helium remains the only storefront creator application.", "Pending, Approved, Rejected and Suspended statuses.", "Customer webhooks, idempotent upsert, tag conflict handling and audit logs.", "Pending applications list/detail, admin notes, status history, validation and success/error feedback."],
     ["Confirm Flow always applies the required tags for the active Helium form.", "Move required Helium fields to an accessible namespace or provide Advanced API access for complete admin-side answers.", "Complete a merchant-observed end-to-end applicant test."]),
    ("Milestone 4 — Creator Dashboard", "PARTIAL",
     ["Approved-access logic and explicit state handling.", "Responsive dashboard layout, profile, collection, recent submissions and activity.", "Cards for sales, earnings, orders, collections and published products.", "Loading state always clears after response completion."],
     ["Add read_orders only after merchant approval.", "Define commission, attribution, refund/cancellation and reporting rules.", "Connect real sales, earnings, orders and top-selling data; remove intentional placeholders."]),
    ("Milestone 5 — Creator Profile", "PARTIAL",
     ["Legal/display name, bio, portfolio and profile-image support.", "Profile image upload with file-type/size controls, persistence and immediate preview.", "Storefront profile and collection links; safe Helium Liquid fallback for logged-in customer."],
     ["Build creator self-service edit/save UI for all profile and social fields.", "Add banner upload and complete public creator profile presentation.", "Complete field accessibility mapping for all Helium answers."]),
    ("Milestone 6 — Collection Management", "PARTIAL",
     ["Approval can create a creator-owned Shopify collection if missing.", "Collection URL and ownership mapping are persisted.", "Published products can be assigned to the correct creator collection."],
     ["Creator collection create/edit UI, banner and thumbnail management.", "Draft/Pending/Published/Archived collection lifecycle and admin review workflow.", "Live collection template and mobile acceptance testing."]),
    ("Milestone 7 — Product Creation Workflow", "PARTIAL",
     ["Base-product selection and creator ownership rules.", "Saved InkyBay design URL validation and design submission persistence.", "Admin approval/rejection states and collection assignment."],
     ["Complete the creator-facing UX for all requested title/description/price/variant choices where merchant policy permits.", "Run full design-save-to-submission-to-published-product live QA."]),
    ("Milestone 8 — Automated Publishing", "MOSTLY COMPLETE",
     ["Admin GraphQL product creation/duplication, draft-first status, creator relationship and collection assignment.", "Idempotency and recovery protect against duplicate products and partial publication.", "Retry-safe failed publishing states and publication-scope handling."],
     ["Confirm final price/image/variant policy.", "Run controlled production approval and verify product page, collection placement, InkyBay configuration, cart and checkout."]),
    ("Milestone 9 — Referral System", "TODO",
     [],
     ["Creator referral links and unique codes.", "Click, order and revenue attribution.", "Creator referral dashboard and admin referral reporting.", "Privacy/consent and attribution-window rules."]),
    ("Milestone 10 — Admin Management", "PARTIAL",
     ["Creator, application and design-submission administration.", "Approve/reject/suspend/reinstate actions, notes and audit trail.", "Sync Existing Creators dry run and confirmed import.", "Setup guide and production-safe errors."],
     ["Collection review, commission rates, earnings and payout management.", "Admin financial reconciliation and export/reporting.", "Final QA, documentation, training and client sign-off."]),
]
for title, status, done, remaining in details:
    doc.add_heading(title, level=2)
    p = doc.add_paragraph()
    r = p.add_run(status)
    r.bold = True
    r.font.color.rgb = RGBColor.from_string(
        GREEN if status in ("DONE", "MOSTLY COMPLETE") else AMBER if status == "PARTIAL" else RED
    )
    if done:
        doc.add_heading("Implemented", level=3)
        add_bullets(doc, done)
    doc.add_heading("Remaining / acceptance work", level=3)
    add_bullets(doc, remaining)

page_break(doc)
doc.add_heading("6. Helium Customer Fields — Required Configuration", level=1)
doc.add_paragraph(
    "Helium is the agreed application interface. The Custom House app is the authority for creator status, approval decisions, collections, dashboard access and submissions."
)
add_table(
    doc,
    ["Configuration", "Required value / behavior"],
    [
        ("Active Helium form", "Form ID lXteLY"),
        ("Applicant tags", "creator-applicant and creator-pending"),
        ("Decision tags", "creator-approved, creator-rejected, creator-suspended"),
        ("Flow trigger", "Customer/form submission for the active Helium form"),
        ("Upsert key", "Shop + Shopify customer GID"),
        ("Status authority", "Do not overwrite a newer app-managed decision with an older customer tag"),
        ("Data minimization", "Do not persist email, phone, address or unnecessary protected customer data"),
    ],
    [1.6, 4.9],
)
doc.add_heading("Known Helium data constraint", level=2)
add_note(
    doc,
    "Action required",
    "The Helium namespace app--960624--helium is reserved to Helium. The Custom House server cannot reliably read all fields from it. For complete admin/profile synchronization, either recreate the required fields under an accessible namespace such as customhouse_creator, or provide Helium Advanced API access. Until then, the logged-in storefront can use safe Liquid fallback values, while standard Shopify customer name data can backfill admin names.",
    LIGHT_AMBER,
    AMBER,
)
doc.add_heading("Mapped field families", level=2)
add_bullets(doc, [
    "Legal name and creator display name",
    "Short biography",
    "Portfolio/social URL and platform-specific social links",
    "Creator profile photo",
    "City, country, terms agreement and application answers",
])

doc.add_heading("7. Quality and Deployment Evidence", level=1)
add_table(
    doc,
    ["Check", "Result", "Notes"],
    [
        ("Lint", "PASS", "ESLint validation passed after the latest creator/dashboard changes."),
        ("Typecheck", "PASS", "React Router type generation and TypeScript no-emit check passed."),
        ("Automated tests", "PASS", "33/33 tests passed across creator application, dashboard, domain and Helium sync suites."),
        ("Production build", "PASS", "Compiled React Router production build passed."),
        ("Database", "PASS", "PostgreSQL provider and reviewed migrations; production deploy uses prisma migrate deploy."),
        ("Health endpoint", "PASS", "Production health check returned status ok."),
        ("Live storefront E2E", "VERIFY", "Final theme, InkyBay, cart, checkout, orders and mobile acceptance remain merchant-controlled."),
    ],
    [1.35, 0.8, 4.35],
)
doc.add_heading("Automated coverage includes", level=2)
add_bullets(doc, [
    "Logged-out, missing, pending, approved, rejected and suspended dashboard states.",
    "Loading-state cleanup and safe API-error handling.",
    "Pending and approved Helium imports, duplicate webhook delivery and conflicting tags.",
    "Lazy dashboard synchronization and app-managed status authority.",
    "Authorization, ownership, allowed URLs, status transitions, idempotency and publishing retry safety.",
])

page_break(doc)
doc.add_heading("8. Remaining Work — Prioritized Plan", level=1)
priorities = [
    ("P0", "Merchant acceptance test", "Create a brand-new Helium applicant; verify pending import, review, approval, tag sync, collection, dashboard, design submission and publishing."),
    ("P0", "Helium field access", "Choose accessible custom metafield namespace or Helium Advanced API; confirm the active Flow and tags."),
    ("P0", "Storefront completion QA", "Review every required page, responsive layout, footer, links, policies, sample products, cart and checkout."),
    ("P1", "Dashboard commerce analytics", "Approve read_orders, define commission/attribution/refund rules, then connect sales, earnings, orders and top sellers."),
    ("P1", "Creator self-service", "Complete profile editing, social fields, banner, and collection create/edit controls."),
    ("P1", "Admin finance", "Commission-rate, earnings, payout and reconciliation management."),
    ("P2", "Referral system", "Referral codes/links, clicks, attributed orders/revenue, creator view and admin reporting."),
    ("P2", "Handover", "Final QA report, operating guide, admin training, rollback instructions and client sign-off."),
]
add_table(doc, ["Priority", "Workstream", "Definition of done"], priorities, [0.6, 1.65, 4.25])

doc.add_heading("9. Merchant-Owned Actions", level=1)
add_bullets(doc, [
    "Approve any new Shopify scopes and reinstall/update the custom-distribution app when requested.",
    "Verify production application URL, auth callback, App Proxy URL and managed-install configuration.",
    "Confirm product metafield definitions, publication selection and Global Product tagging.",
    "Publish/enable the latest theme extension blocks and preserve the live header changes.",
    "Confirm Helium form/Flow tags and resolve field namespace or API access.",
    "Configure and verify InkyBay products, DesignLab options, prices, cart behavior and checkout.",
    "Approve commission, attribution, refunds, payout and reporting policies before financial dashboards are connected.",
])
add_note(
    doc,
    "Release rule",
    "Do not bulk-change products, remove Helium/Flow, or change merchant-owned InkyBay configuration without explicit merchant confirmation.",
    LIGHT_RED,
    RED,
)

doc.add_heading("10. Recommended Client Message", level=1)
doc.add_paragraph(
    "The Custom House Creator Marketplace now has a stable production foundation and a working core creator lifecycle: Helium application intake, synchronization, admin review, approval statuses, dashboard access, design submission and safe draft publishing. The next phase is not a rebuild. It is focused completion: resolving full Helium field access, connecting approved order/commission data, adding creator self-service profile and collection editing, implementing referrals and payouts, and completing merchant-led storefront and checkout acceptance testing."
)

page_break(doc)
doc.add_heading("Appendix A — Full Checklist Status", level=1)
appendix = {
    "Milestone 1 — Storefront Foundation": [
        ("Shopify store setup", "DONE"), ("Theme setup and activation", "DONE"), ("Announcement bar", "VERIFY"),
        ("Header layout and logo", "DONE"), ("Main navigation menu", "DONE"), ("Footer layout and links", "VERIFY"),
        ("Homepage hero banner", "VERIFY"), ("Homepage featured collections", "VERIFY"), ("Homepage featured products", "VERIFY"),
        ("Homepage content sections", "VERIFY"), ("Responsive homepage layout", "VERIFY"), ("Collection structure", "PARTIAL"),
        ("Collection page template", "VERIFY"), ("Product page template", "PARTIAL"), ("Cart page", "VERIFY"), ("Checkout setup", "VERIFY"),
        ("About Us page", "VERIFY"), ("Contact page/form", "VERIFY"), ("Blog setup", "VERIFY"), ("FAQ page", "VERIFY"),
        ("Legal/policy pages", "VERIFY"), ("Sample products", "VERIFY"), ("Sample images/content", "VERIFY"),
        ("Broken-link check", "VERIFY"), ("Basic speed/UX review", "VERIFY"), ("Client design approval", "TODO"),
    ],
    "Milestone 2 — InkyBay Product Personalizer": [
        ("InkyBay installed", "DONE"), ("Plan and initial settings", "VERIFY"), ("DesignLab configured", "VERIFY"),
        ("Editable canvas", "VERIFY"), ("Artwork upload", "VERIFY"), ("Text/logo tools", "VERIFY"),
        ("Customization fields", "VERIFY"), ("Product connection", "PARTIAL"), ("Pricing/add-ons", "VERIFY"),
        ("Mobile/responsive DesignLab", "VERIFY"), ("Cart customization display", "VERIFY"),
        ("Checkout/customization flow", "VERIFY"), ("Test customization order", "VERIFY"),
    ],
    "Milestone 3 — Creator Onboarding / Application": [
        ("Backend stack confirmed", "DONE"), ("Database setup", "DONE"), ("Application data structure", "DONE"),
        ("Application form/page via Helium", "DONE"), ("Social profile fields", "PARTIAL"), ("Audience information fields", "PARTIAL"),
        ("Application submit/sync API", "DONE"), ("Application saved in database", "DONE"), ("Default Pending status", "DONE"),
        ("Pending/Approved/Rejected/Suspended statuses", "DONE"), ("Admin application list", "DONE"),
        ("Admin application detail", "DONE"), ("Approve/Reject/Suspend actions", "DONE"), ("Admin notes", "DONE"),
        ("Status/history tracking", "DONE"), ("Validation", "DONE"), ("Success/error messages", "DONE"),
    ],
    "Milestone 4 — Creator Dashboard": [
        ("Approved creator access", "DONE"), ("Dashboard route/page", "DONE"), ("Overview layout", "DONE"),
        ("Total sales card", "PARTIAL"), ("Total earnings card", "PARTIAL"), ("Orders count card", "PARTIAL"),
        ("Collections count card", "DONE"), ("Top selling products", "PARTIAL"), ("Recent activity", "DONE"),
        ("Dashboard data connection/placeholder", "DONE"), ("Mobile responsive dashboard", "DONE"),
    ],
    "Milestone 5 — Creator Profile": [
        ("Profile page structure", "PARTIAL"), ("Profile photo upload", "DONE"), ("Banner image upload", "TODO"),
        ("Bio section", "DONE"), ("Social fields", "PARTIAL"), ("Edit profile", "TODO"), ("Save profile changes", "PARTIAL"),
        ("Public creator profile", "PARTIAL"), ("Profile link structure", "DONE"), ("Responsive layout", "PARTIAL"),
    ],
    "Milestone 6 — Collection Management": [
        ("Creator collection data structure", "PARTIAL"), ("Create collection page", "TODO"), ("Collection title/description", "PARTIAL"),
        ("Banner upload", "TODO"), ("Thumbnail upload", "TODO"), ("Edit collection", "TODO"),
        ("Draft/Pending/Published/Archived statuses", "TODO"), ("Admin review", "PARTIAL"),
        ("Storefront collection template", "PARTIAL"), ("Responsive collection page", "VERIFY"),
    ],
    "Milestone 7 — Product Creation Workflow": [
        ("Create product action", "PARTIAL"), ("Select base product", "DONE"), ("Connect InkyBay DesignLab", "DONE"),
        ("Save creator design", "DONE"), ("Assign collection", "DONE"), ("Title/description", "PARTIAL"),
        ("Price", "PARTIAL"), ("Variants", "PARTIAL"), ("Images/previews", "PARTIAL"), ("Save as Draft", "DONE"),
        ("Submit for Review", "DONE"), ("Responsive product flow", "PARTIAL"),
    ],
    "Milestone 8 — Automated Publishing": [
        ("Shopify Admin API connection", "DONE"), ("Duplicate/create base product", "DONE"), ("Replace images", "PARTIAL"),
        ("Update title/description", "DONE"), ("Price setup", "PARTIAL"), ("Duplicate variants", "DONE"),
        ("Creator-product relation", "DONE"), ("Collection assignment", "DONE"), ("Publish storefront", "DONE"),
        ("Display in creator collection", "DONE"), ("End-to-end testing", "VERIFY"),
    ],
    "Milestone 9 — Referral System": [
        ("Referral-link structure", "TODO"), ("Unique code", "TODO"), ("Click tracking", "TODO"), ("Order tracking", "TODO"),
        ("Revenue tracking", "TODO"), ("Creator dashboard", "TODO"), ("Admin tracking", "TODO"),
    ],
    "Milestone 10 — Admin Management": [
        ("Creator list/status/details", "DONE"), ("Creator actions", "DONE"), ("Application review", "DONE"),
        ("Collection review", "TODO"), ("Product submission review", "DONE"), ("Commission rates", "TODO"),
        ("Earnings", "TODO"), ("Payout tracking", "TODO"), ("Notes/history", "DONE"), ("Final QA", "PARTIAL"),
    ],
    "Overall Final QA": [
        ("Application flow", "PARTIAL"), ("Creator approval flow", "PARTIAL"), ("Dashboard data", "PARTIAL"),
        ("Profile update flow", "PARTIAL"), ("Collection flow", "PARTIAL"), ("Product creation flow", "PARTIAL"),
        ("InkyBay customization", "VERIFY"), ("Product publishing", "PARTIAL"), ("Cart/checkout", "VERIFY"),
        ("Test order", "VERIFY"), ("Admin controls", "PARTIAL"), ("Referral tracking", "TODO"),
        ("Mobile/responsive", "PARTIAL"), ("Broken links", "VERIFY"), ("Console errors", "VERIFY"),
        ("No placeholders/dummy data", "TODO"), ("Documentation", "DONE"), ("Client handover", "TODO"),
    ],
}
for milestone, items in appendix.items():
    doc.add_heading(milestone, level=2)
    rows = [(status, item) for item, status in items]
    add_table(doc, ["Status", "Checklist item"], rows, [1.25, 5.25], 0)

doc.add_heading("Appendix B — Technical Inventory", level=1)
add_table(
    doc,
    ["Category", "Current implementation"],
    [
        ("Core stack", "Shopify React Router, TypeScript, App Bridge, Admin GraphQL"),
        ("Database", "Prisma + PostgreSQL"),
        ("Core models", "Session, ShopConfig, Creator, CreatorApplication, DesignSubmission, AuditLog"),
        ("Admin routes", "/app, /app/applications, /app/creators, /app/design-submissions, /app/products, /app/settings, /app/setup-guide"),
        ("Proxy APIs", "Dashboard, creator profile upload, applications and design submissions"),
        ("Webhooks", "Customer create/update, product/collection delete, uninstall/scopes and compliance topics"),
        ("Deployment", "Render production app with compiled React Router start command"),
        ("Latest verified code line", "Creator identity, Helium profile fallback, admin review, dashboard overview and avatar fixes"),
    ],
    [1.4, 5.1],
)
doc.add_paragraph(
    "This report intentionally excludes customer personal information, access tokens, cookies, database URLs and other secrets."
).italic = True

# Document metadata and save
props = doc.core_properties
props.title = "Custom House Creator Marketplace — Project Progress Report"
props.subject = "Client status, milestone completion, remaining roadmap and deployment readiness"
props.author = "Custom House Engineering"
props.keywords = "Shopify, Creator Marketplace, Helium, InkyBay, progress report"
OUT.parent.mkdir(parents=True, exist_ok=True)
doc.save(OUT)
print(OUT)
